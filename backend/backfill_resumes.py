import asyncio, uuid, os
from motor.motor_asyncio import AsyncIOMotorClient

UPLOAD_DIR = os.path.abspath('uploads')

def make_docx(rd, name, email, score):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    rd = rd or {}
    def safe(v): return str(v) if v else ''
    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(safe(name))
    run.bold = True
    doc.add_paragraph(safe(email or '')).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('_' * 60)
    total = (score or {}).get('total_score', 0) or 0
    skill = (score or {}).get('skill_score', 0) or 0
    exp   = (score or {}).get('experience_score', 0) or 0
    sp = doc.add_paragraph()
    sp.add_run(f"ATS Score: {total:.1f}%  |  Skills: {skill:.1f}%  |  Experience: {exp:.1f}%").bold = True
    doc.add_paragraph('_' * 60)
    if rd.get('summary'):
        doc.add_heading('Professional Summary', 2)
        doc.add_paragraph(rd['summary'])
    if rd.get('skills'):
        doc.add_heading('Skills', 2)
        skills = rd['skills']
        doc.add_paragraph(' | '.join(skills) if isinstance(skills, list) else str(skills))
    experience = rd.get('experience', [])
    if experience:
        doc.add_heading('Work Experience', 2)
        if isinstance(experience, list):
            for ex in experience:
                if isinstance(ex, dict):
                    line = ex.get('title', '')
                    if ex.get('company'): line += ' -- ' + ex['company']
                    p = doc.add_paragraph()
                    p.add_run(line).bold = True
                    if ex.get('description'): doc.add_paragraph(str(ex['description']))
    education = rd.get('education', [])
    edu_level = rd.get('education_level', '')
    if education or edu_level:
        doc.add_heading('Education', 2)
        if isinstance(education, list):
            for edu in education:
                if isinstance(edu, dict):
                    line = edu.get('degree', '')
                    if edu.get('institution'): line += ' -- ' + edu['institution']
                    if edu.get('year'): line += ' (' + str(edu['year']) + ')'
                    p = doc.add_paragraph()
                    p.add_run(line).bold = True
        if edu_level: doc.add_paragraph('Highest Qualification: ' + str(edu_level))
    fname = str(uuid.uuid4()) + '.docx'
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    doc.save(os.path.join(UPLOAD_DIR, fname))
    return fname

async def backfill():
    c = AsyncIOMotorClient('mongodb://localhost:27017')
    db = c.resume_screening
    records = await db.resumes.find({'resume_filename': None}).to_list(100)
    print(f'Found {len(records)} records with no resume file')
    for r in records:
        score = r.get('score', {})
        total = (score or {}).get('total_score', 0) or 0
        if total >= 70:
            fname = make_docx(r.get('resume_data'), r.get('candidate_name', ''), r.get('candidate_email', ''), score)
            await db.resumes.update_one({'_id': r['_id']}, {'$set': {'resume_filename': fname}})
            print(f"  Backfilled: {r.get('candidate_name')} -> {fname}")
        else:
            print(f"  Skipped (score={total:.1f}): {r.get('candidate_name')}")
    print("Done.")

asyncio.run(backfill())
