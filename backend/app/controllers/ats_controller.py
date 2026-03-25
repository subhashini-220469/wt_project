import asyncio
from typing import List
from fastapi import UploadFile, HTTPException
from ..services.extraction import DocumentExtractor, LLMParser
from ..services.scoring import ScoringEngine
from ..db.database import db
from ..models.models import ResumeData, JDData
import os
import uuid

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "uploads")


def _generate_resume_docx(resume_data_dict: dict, candidate_name: str, candidate_email: str, score_dict: dict) -> str:
    """Generate a .docx resume from parsed data and save it to uploads/. Returns saved filename."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO

        rd = resume_data_dict or {}

        def safe(val):
            return str(val) if val is not None else ""

        doc = Document()

        # Header
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(safe(candidate_name))
        run.bold = True
        run.font.size = Pt(20)

        contact_parts = [candidate_email] if candidate_email else []
        if rd.get("phone"):
            contact_parts.append(rd["phone"])
        if rd.get("location"):
            contact_parts.append(rd["location"])
        if contact_parts:
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub.add_run(safe(" | ".join(contact_parts)))

        doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ATS Score
        total_score = (score_dict or {}).get("total_score", 0) or 0
        skill_score = (score_dict or {}).get("skill_score", 0) or 0
        exp_score   = (score_dict or {}).get("experience_score", 0) or 0
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = sp.add_run(f"ATS Score: {total_score:.1f}%  |  Skills: {skill_score:.1f}%  |  Experience: {exp_score:.1f}%")
        r1.bold = True

        doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        def section(heading):
            h = doc.add_heading(heading, level=2)

        # Summary
        summary = rd.get("summary", "")
        if summary:
            section("Professional Summary")
            doc.add_paragraph(summary)

        # Skills
        skills = rd.get("skills", [])
        if skills:
            section("Skills")
            doc.add_paragraph(" | ".join(skills) if isinstance(skills, list) else str(skills))

        # Experience
        experience = rd.get("experience", [])
        if experience:
            section("Work Experience")
            if isinstance(experience, list):
                for exp in experience:
                    if isinstance(exp, dict):
                        line = exp.get("title", "")
                        if exp.get("company"): line += f" -- {exp['company']}"
                        if exp.get("duration") or exp.get("years"): line += f" ({exp.get('duration') or exp.get('years', '')})"
                        p = doc.add_paragraph()
                        p.add_run(safe(line)).bold = True
                        if exp.get("description"): doc.add_paragraph(safe(str(exp["description"])))

        # Education
        education = rd.get("education", [])
        edu_level = rd.get("education_level", "")
        if education or edu_level:
            section("Education")
            if isinstance(education, list):
                for edu in education:
                    if isinstance(edu, dict):
                        line = edu.get("degree", "")
                        if edu.get("institution"): line += f" -- {edu['institution']}"
                        if edu.get("year"): line += f" ({edu['year']})"
                        p = doc.add_paragraph()
                        p.add_run(safe(line)).bold = True
            if edu_level:
                doc.add_paragraph(f"Highest Qualification: {edu_level}")

        # Certifications
        certs = rd.get("certifications", [])
        if certs:
            section("Certifications")
            for c in (certs if isinstance(certs, list) else [str(certs)]):
                doc.add_paragraph(f"- {c}")

        # Save to uploads/
        if not os.path.exists(UPLOAD_DIR):
            os.makedirs(UPLOAD_DIR)
        filename = f"{uuid.uuid4()}.docx"
        save_path = os.path.join(UPLOAD_DIR, filename)
        doc.save(save_path)
        print(f"Generated resume docx: {filename}")
        return filename
    except Exception as e:
        print(f"[WARN] Could not generate resume docx: {e}")
        return None

class ATSController:
    @staticmethod
    def _save_file(file: UploadFile) -> str:
        if not os.path.exists(UPLOAD_DIR):
            os.makedirs(UPLOAD_DIR)
        
        # Generate a unique filename to avoid collisions
        ext = os.path.splitext(file.filename)[1]
        unique_name = f"{uuid.uuid4()}{ext}"
        save_path = os.path.join(UPLOAD_DIR, unique_name)
        
        # We need to read the content to save it, but we might have already read it
        # Actually, it's safer to read it once and reuse it if needed
        return unique_name, save_path
    @staticmethod
    async def parse_resume(file: UploadFile):
        try:
            content = await file.read()
            
            # Save the file
            unique_name, save_path = ATSController._save_file(file)
            with open(save_path, "wb") as f:
                f.write(content)

            if file.filename.endswith(".pdf"):
                text = DocumentExtractor.extract_text_from_pdf(content)
            elif file.filename.endswith(".docx"):
                text = DocumentExtractor.extract_text_from_docx(content)
            else:
                raise HTTPException(status_code=400, detail="Unsupported format")
            
            resume_dict = await LLMParser.parse_resume(text)
            resume_data = ResumeData(**resume_dict)
            
            return {
                "name": resume_data.name,
                "resume_data": resume_data.model_dump(),
                "resume_filename": unique_name
            }
        except Exception as e:
            print(f"Error parsing resume: {e}")
            raise HTTPException(status_code=500, detail=str(e))

    @staticmethod
    async def get_profile(email: str):
        try:
            profile = await db.db.profiles.find_one({"email": email})
            if profile:
                profile["_id"] = str(profile["_id"])
                return profile
            return None
        except Exception as e:
            print(f"Error fetching profile: {e}")
            raise HTTPException(status_code=500, detail=str(e))

    @staticmethod
    async def upsert_profile(email: str, profile_data: dict):
        try:
            # Ensure email is in the data
            profile_data["email"] = email
            
            # Use email as unique identifier for profiles
            result = await db.db.profiles.update_one(
                {"email": email},
                {"$set": profile_data},
                upsert=True
            )
            return {"message": "Profile updated successfully"}
        except Exception as e:
            print(f"Error updating profile: {e}")
            raise HTTPException(status_code=500, detail=str(e))

    @staticmethod
    async def apply_to_job(job_id: str, name: str, email: str, file: UploadFile, screening_answers: dict, resume_data_override: dict = None):
        # 1. Fetch Job from DB to get structured_jd
        from bson import ObjectId
        job = await db.db.jobs.find_one({"_id": ObjectId(job_id)})
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        
        jd_data_dict = job.get("structured_jd")
        if not jd_data_dict:
            jd_data_dict = await LLMParser.parse_jd(job["description"])
        
        jd_data = JDData(**jd_data_dict)

        # 2a. Check for duplicate application
        existing_app = await db.db.resumes.find_one({"jd_id": job_id, "candidate_email": email})
        if existing_app:
            raise HTTPException(status_code=400, detail="You have already applied for this position.")

        # 2. Get Resume Data (either from file or from override)
        resume_filename = None
        if resume_data_override:
            # resume_data_override might contain resume_filename if it was parsed earlier
            resume_filename = resume_data_override.get("resume_filename")
            resume_data = ResumeData(**resume_data_override.get("resume_data", resume_data_override))
        else:
            content = await file.read()
            
            # Save the file
            unique_name, save_path = ATSController._save_file(file)
            with open(save_path, "wb") as f:
                f.write(content)
            resume_filename = unique_name

            if file.filename.endswith(".pdf"):
                text = DocumentExtractor.extract_text_from_pdf(content)
            elif file.filename.endswith(".docx"):
                text = DocumentExtractor.extract_text_from_docx(content)
            else:
                raise HTTPException(status_code=400, detail="Unsupported format")
            
            resume_dict = await LLMParser.parse_resume(text)
            resume_data = ResumeData(**resume_dict)

        # 3. Final polish on resume data
        resume_data.name = name
        resume_data.email = email
        
        # 4. Score
        score_res = ScoringEngine.score_resume(resume_data, jd_data)

        # 4b. If score >= 70 and no uploaded file, auto-generate a .docx and save it
        if score_res.total_score >= 70 and not resume_filename:
            resume_filename = _generate_resume_docx(
                resume_data.model_dump(), name, email, score_res.model_dump()
            )

        # 5. Store in Resumes collection
        from datetime import datetime
        resume_record = {
            "jd_id": job_id,
            "job_title": jd_data.job_title,
            "filename": file.filename if file else "pre-parsed",
            "resume_filename": resume_filename,  # Always set now for score >= 70
            "candidate_name": name,
            "candidate_email": email,
            "screening_answers": screening_answers,
            "resume_data": resume_data.model_dump(),
            "score": score_res.model_dump(),
            "status": "applied",
            "applied_at": datetime.utcnow().isoformat()
        }

        await db.db.resumes.insert_one(resume_record)
        
        # 6. Notify the Recruiter
        try:
            hr_id_str = job.get("posted_by")
            if hr_id_str:
                hr_user = await db.users_db.users.find_one({"_id": ObjectId(hr_id_str)})
                if hr_user and "email" in hr_user:
                    hr_notif = {
                        "user_email": hr_user["email"],
                        "message": f"📈 New application for {jd_data.job_title} from {name}",
                        "is_read": False,
                        "type": "new_app",
                        "created_at": datetime.utcnow().isoformat(),
                        "link": "/app/analytics" # Or dashboard
                    }
                    await db.db.notifications.insert_one(hr_notif)
                    print(f"Notified HR ({hr_user['email']}) about new application.")
        except Exception as e:
            print(f"Failed to notify HR: {e}")

        print(f"Successfully stored application for {name} in 'resumes' collection.")
        
        return {
            "message": "Application submitted successfully",
            "score": score_res.total_score,
            "details": score_res.model_dump()
        }
