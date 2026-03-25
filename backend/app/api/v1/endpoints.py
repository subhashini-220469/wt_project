from fastapi import APIRouter, UploadFile, File, Form, HTTPException, BackgroundTasks
from typing import List, Optional
import datetime
import json
from ...controllers.ats_controller import ATSController
from ...db.database import db
from ...services.mailer import Mailer
from ...services.scoring import ScoringEngine
from ...services.extraction import LLMParser
from ...models.models import ResumeData, JDData, ScoringResult, JobPosting, Notification
from bson import ObjectId
from pydantic import BaseModel
import datetime
from fastapi.responses import FileResponse
import os

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), "uploads")

router = APIRouter()

async def process_job_requirements(job_id: str, description: str):
    """Background task to extract structured requirements from JD using AI."""
    try:
        print(f"Background: Starting JD parsing for Job {job_id}...")
        structured_dict = await LLMParser.parse_jd(description)
        # Validate data
        validated_jd = JDData(**structured_dict)
        
        # update the job record
        await db.db.jobs.update_one(
            {"_id": ObjectId(job_id)},
            {"$set": {"structured_jd": validated_jd.model_dump()}}
        )
        print(f"Background: JD parsing complete for Job {job_id}.")
    except Exception as e:
        print(f"Error in background JD parsing for {job_id}: {e}")


@router.post("/jobs")
async def create_job(job_post: JobPosting, background_tasks: BackgroundTasks):
    try:
        # Save the job first to get an ID
        job_dict = job_post.model_dump()
        result = await db.db.jobs.insert_one(job_dict)
        job_id = result.inserted_id
        
        # Dispatch the slow AI parsing to a background task
        background_tasks.add_task(process_job_requirements, str(job_id), job_post.description)
        
        # Create notifications for all candidates ('user' role)
        try:
            # Find all users with role 'user' from the node backend's collection
            candidate_cursor = db.users_db.users.find({"role": "user"})
            candidates = await candidate_cursor.to_list(length=1000)
            
            notifications = []
            for candidate in candidates:
                notif = {
                    "user_email": candidate["email"],
                    "message": f"🔔 HR added a new job: {job_post.job_title} at {job_post.company}",
                    "is_read": False,
                    "type": "new_job",
                    "created_at": datetime.datetime.utcnow().isoformat(),
                    "link": "/app",
                    "job_id": str(job_id)
                }
                notifications.append(notif)
            
            if notifications:
                await db.db.notifications.insert_many(notifications)
        except Exception as e:
            print(f"Error creating new job notifications: {e}")

        return {
            "id": str(job_id), 
            "message": "Job posted successfully. Requirements are being analyzed in the background."
        }
    except Exception as e:
        print(f"Error creating job: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/jobs")
async def get_jobs(posted_by: Optional[str] = None):
    query = {}
    if posted_by:
        query["posted_by"] = posted_by
    cursor = db.db.jobs.find(query)
    jobs = await cursor.to_list(length=100)
    
    now = datetime.datetime.utcnow()
    for j in jobs:
        j["_id"] = str(j["_id"])
        
        # Check if job deadline has passed
        if j.get("status") == "open" and "deadline" in j and j["deadline"]:
            try:
                deadline_dt = datetime.datetime.strptime(str(j["deadline"]), "%Y-%m-%d")
                expiry_dt = deadline_dt + datetime.timedelta(days=1)
                if now >= expiry_dt:
                    j["status"] = "expired"
            except Exception:
                pass
    return jobs


@router.get("/results/{jd_id}")
async def get_results(jd_id: str):
    cursor = db.db.resumes.find({"jd_id": jd_id})
    resumes = await cursor.to_list(length=100)
    for r in resumes:
        r["_id"] = str(r["_id"])
    
    # Sort by score
    resumes.sort(key=lambda x: x["score"]["total_score"], reverse=True)
    return resumes

@router.get("/my-applications/{email}")
async def get_my_applications(email: str):
    try:
        cursor = db.db.resumes.find({"candidate_email": email})
        apps = await cursor.to_list(length=50)
        for a in apps:
            a["_id"] = str(a["_id"])
        return apps
    except Exception as e:
        print(f"Error fetching personal apps: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/jds")
async def get_all_jds(posted_by: Optional[str] = None):
    # Keep this for backward compatibility or general list
    query = {}
    if posted_by:
        query["posted_by"] = posted_by
    cursor = db.db.jobs.find(query)
    jds = await cursor.to_list(length=100)
    for j in jds:
        id_obj = ObjectId(j["_id"])
        j["_id"] = str(id_obj)
        j["created_at"] = id_obj.generation_time.isoformat()
    return jds


@router.post("/parse-resume")
async def parse_resume(file: UploadFile = File(...)):
    return await ATSController.parse_resume(file)

@router.get("/profiles/{email}")
async def get_profile(email: str):
    profile = await ATSController.get_profile(email)
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")
    return profile

@router.post("/profiles/{email}")
async def upsert_profile(email: str, profile_data: dict):
    return await ATSController.upsert_profile(email, profile_data)

@router.post("/apply/{job_id}")
async def apply_to_job(
    job_id: str,
    name: str = Form(...),
    email: str = Form(...),
    resume: UploadFile = File(None), # Optional if override provided
    screening_answers: str = Form("{}"),
    resume_data_override: str = Form(None) # JSON string
):
    try:
        import json
        answers_dict = json.loads(screening_answers)
        override_dict = json.loads(resume_data_override) if resume_data_override else None
        return await ATSController.apply_to_job(job_id, name, email, resume, answers_dict, override_dict)
    except Exception as e:
        print(f"Error in /apply: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/rescore/{job_id}")
async def rescore_job(job_id: str, resume_data: ResumeData):
    try:
        job = await db.db.jobs.find_one({"_id": ObjectId(job_id)})
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        
        jd_dict = job.get("structured_jd")
        if not jd_dict:
            # Fallback if JD wasn't parsed (e.g. old data)
            jd_dict = await LLMParser.parse_jd(job.get("description", ""))

        jd_data = JDData(**jd_dict)
        score_res = ScoringEngine.score_resume(resume_data, jd_data)
        return score_res
    except Exception as e:
        print(f"Error in /rescore: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/jobs/{job_id}")
async def delete_job(job_id: str):
    try:
        result = await db.db.jobs.delete_one({"_id": ObjectId(job_id)})
        if result.deleted_count == 0:
            raise HTTPException(status_code=404, detail="Job not found")
        # Also cleanup resumes for this job
        await db.db.resumes.delete_many({"jd_id": job_id})
        return {"message": "Job and associated applications deleted successfully"}
    except Exception as e:
        print(f"Error deleting job: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.patch("/jobs/{job_id}/status")
async def update_job_status(job_id: str, status: str = Form(...)):
    try:
        result = await db.db.jobs.update_one(
            {"_id": ObjectId(job_id)},
            {"$set": {"status": status}}
        )
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Job not found")
        return {"message": f"Job status updated to {status}"}
    except Exception as e:
        print(f"Error updating job status: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class EmailRequest(BaseModel):
    jd_id: str
    recipient_emails: List[str]
    subject: str
    body: str

@router.get("/analytics/jobs")
async def get_job_analytics(posted_by: Optional[str] = None):
    try:
        # Get jobs filtered by HR
        query = {}
        if posted_by:
            query["posted_by"] = posted_by
            
        cursor = db.db.jobs.find(query)
        jobs = await cursor.to_list(length=100)
        
        analytics = []
        for job in jobs:
            job_id = str(job["_id"])
            
            # Count applications
            total_apps = await db.db.resumes.count_documents({"jd_id": job_id})
            
            # Count selected (Shortlisted score >= 70)
            # Find all matching jd_id
            cursor_res = db.db.resumes.find({"jd_id": job_id})
            resumes = await cursor_res.to_list(length=1000)
            
            selected = 0
            interviewed = 0
            for r in resumes:
                score = r.get("score", {}).get("total_score", 0)
                if score >= 70:
                    selected += 1
                
                # Check status for 'interviewed'
                if r.get("status") == "interviewed":
                    interviewed += 1
            
            analytics.append({
                "job_id": job_id,
                "job_title": job.get("job_title", "Untitled"),
                "company": job.get("company", "Unknown"),
                "total_applied": total_apps,
                "selected": selected,
                "interviews_done": interviewed,
                "status": job.get("status", "open")
            })
            
        return analytics
    except Exception as e:
        print(f"Error in analytics: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/send-emails")
async def send_bulk_emails(req: EmailRequest, background_tasks: BackgroundTasks):
    if not req.recipient_emails:
        raise HTTPException(status_code=400, detail="No recipients provided")
    
    # Update status to 'interviewed' for these candidates
    await db.db.resumes.update_many(
        {"jd_id": req.jd_id, "resume_data.email": {"$in": req.recipient_emails}},
        {"$set": {"status": "interviewed"}}
    )
    
    # Create notifications for these candidates
    try:
        job = await db.db.jobs.find_one({"_id": ObjectId(req.jd_id)})
        job_title = job.get("job_title", "Position") if job else "Position"
        
        candidate_notifs = []
        for email in req.recipient_emails:
            notif = {
                "user_email": email,
                "message": f"📨 You have a new message regarding your application for {job_title}",
                "is_read": False,
                "type": "interview",
                "created_at": datetime.datetime.utcnow().isoformat(),
                "link": "/app/my-apps"
            }
            candidate_notifs.append(notif)
        
        if candidate_notifs:
            await db.db.notifications.insert_many(candidate_notifs)
    except Exception as e:
        print(f"Error creating email notifications: {e}")
    
    # Use FastAPI BackgroundTasks for queueing
    background_tasks.add_task(
        Mailer.send_bulk_emails, 
        req.recipient_emails, 
        req.subject, 
        req.body
    )
    
    return {"message": f"Bulk email task started for {len(req.recipient_emails)} candidates."}

@router.get("/notifications")
async def get_notifications(email: str):
    try:
        # Before returning notifications, check for upcoming deadlines (e.g. tomorrow)
        # to auto-generate reminders for this user
        # Check user role - HRs should NOT get deadline reminders
        user = await db.users_db.users.find_one({"email": email})
        is_candidate = user and user.get("role") == "user"
        
        if is_candidate:
            now = datetime.datetime.utcnow().date()
            tomorrow = now + datetime.timedelta(days=1)
            tomorrow_str = tomorrow.strftime("%Y-%m-%d")
            
            # Find jobs closing tomorrow
            closing_jobs = await db.db.jobs.find({"deadline": tomorrow_str, "status": "open"}).to_list(length=10)
            
            for job in closing_jobs:
                job_id = str(job["_id"])
                # Check if we already notified this user about this specific deadline
                existing = await db.db.notifications.find_one({
                    "user_email": email, 
                    "job_id": job_id, 
                    "type": "deadline"
                })
                
                if not existing:
                    deadline_notif = {
                        "user_email": email,
                        "message": f"⏳ Last day to apply! {job['job_title']} at {job['company']} closes tomorrow.",
                        "is_read": False,
                        "type": "deadline",
                        "created_at": datetime.datetime.utcnow().isoformat(),
                        "link": "/app/discover",
                        "job_id": job_id
                    }
                    await db.db.notifications.insert_one(deadline_notif)
                    print(f"Generated auto-deadline reminder for {email} regarding job {job_id}")

        # Now fetch all (including new ones)
        cursor = db.db.notifications.find({"user_email": email}).sort("created_at", -1)
        notifs = await cursor.to_list(length=50)
        for n in notifs:
            n["_id"] = str(n["_id"])
        return notifs
    except Exception as e:
        print(f"Error fetching notifications: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.patch("/notifications/{notif_id}/read")
async def mark_notification_read(notif_id: str):
    try:
        result = await db.db.notifications.update_one(
            {"_id": ObjectId(notif_id)},
            {"$set": {"is_read": True}}
        )
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Notification not found")
        return {"message": "Notification marked as read"}
    except Exception as e:
        print(f"Error marking notification read: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/resumes/download/{filename}")
async def download_resume(filename: str):
    file_path = os.path.join(UPLOAD_DIR, filename)
    if os.path.exists(file_path):
        return FileResponse(file_path)
    
    # If the original file is missing from disk (wiped or seeded data), gracefully fallback to DOC generation
    # First, find the application ID associated with this filename
    record = await db.db.resumes.find_one({"resume_filename": filename})
    if record:
        return await generate_resume_doc(str(record["_id"]))
    
    raise HTTPException(status_code=404, detail="Resume file not found")


@router.get("/resumes/generate-doc/{application_id}")
async def generate_resume_doc(application_id: str):
    """Generate a formatted Word Document from the stored parsed resume data."""
    try:
        from io import BytesIO
        from fastapi.responses import StreamingResponse
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        record = await db.db.resumes.find_one({"_id": ObjectId(application_id)})
        if not record:
            raise HTTPException(status_code=404, detail="Application record not found")

        rd = record.get("resume_data", {})
        score = record.get("score", {})
        candidate_name = record.get("candidate_name", rd.get("name", "Candidate")) or "Candidate"
        candidate_email = record.get("candidate_email", rd.get("email", "")) or ""

        def safe(val):
            return str(val) if val is not None else ""

        doc = Document()

        # Header
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(safe(candidate_name))
        run.bold = True
        run.font.size = Pt(22)

        contact_parts = []
        if candidate_email:
            contact_parts.append(candidate_email)
        if rd.get("phone"):
            contact_parts.append(rd["phone"])
        if rd.get("location"):
            contact_parts.append(rd["location"])

        if contact_parts:
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            import docx
            s_run = subtitle.add_run(safe(" | ".join(contact_parts)))
            s_run.font.size = Pt(10)
            s_run.font.color.rgb = docx.shared.RGBColor(100, 116, 139)

        doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ATS Banner
        total_score = score.get("total_score", 0) or 0
        skill_score = score.get("skill_score", 0) or 0
        exp_score   = score.get("experience_score", 0) or 0

        score_p = doc.add_paragraph()
        score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = score_p.add_run(f"ATS Match Score: {total_score:.1f}%  |  ")
        r1.bold = True
        score_p.add_run(f"Skills: {skill_score:.1f}%  |  Experience: {exp_score:.1f}%")

        doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        def add_section(heading):
            import docx
            h = doc.add_heading(heading, level=2)
            h.runs[0].font.color.rgb = docx.shared.RGBColor(30, 41, 59)

        # Summary
        summary = rd.get("summary", "")
        if summary:
            add_section("Professional Summary")
            doc.add_paragraph(summary)

        # Skills
        skills = rd.get("skills", [])
        if skills:
            add_section("Skills")
            skills_text = " • ".join(skills) if isinstance(skills, list) else str(skills)
            doc.add_paragraph(skills_text)

        # Experience
        experience = rd.get("experience", [])
        if experience:
            add_section("Work Experience")
            if isinstance(experience, list):
                for exp in experience:
                    if isinstance(exp, dict):
                        title_line = exp.get("title", "")
                        if exp.get("company"):
                            title_line += f"  --  {exp['company']}"
                        if exp.get("duration") or exp.get("years"):
                            title_line += f"  ({exp.get('duration') or exp.get('years', '')})"
                        p = doc.add_paragraph()
                        p.add_run(safe(title_line)).bold = True
                        if exp.get("description"):
                            doc.add_paragraph(str(exp["description"]))
                    else:
                        doc.add_paragraph(str(exp))
            else:
                doc.add_paragraph(str(experience))

        # Education
        education = rd.get("education", [])
        edu_level  = rd.get("education_level", "")
        if education or edu_level:
            add_section("Education")
            if isinstance(education, list):
                for edu in education:
                    if isinstance(edu, dict):
                        edu_line = edu.get("degree", "")
                        if edu.get("institution"):
                            edu_line += f"  --  {edu['institution']}"
                        if edu.get("year"):
                            edu_line += f"  ({edu['year']})"
                        p = doc.add_paragraph()
                        p.add_run(safe(edu_line)).bold = True
                    else:
                        doc.add_paragraph(str(edu))
            elif education:
                doc.add_paragraph(str(education))
            if edu_level:
                doc.add_paragraph(f"Highest Qualification: {edu_level}")

        # Certifications
        certs = rd.get("certifications", [])
        if certs:
            add_section("Certifications")
            if isinstance(certs, list):
                for c in certs:
                    doc.add_paragraph(str(c))
            else:
                doc.add_paragraph(str(certs))

        # Produce Bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        import urllib.parse
        safe_name = urllib.parse.quote(safe(candidate_name).replace(" ", "_"))
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{safe_name}_Analysis.docx"'
            }
        )
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        raise HTTPException(status_code=500, detail=str(e) + "\n" + tb)
