import asyncio, os, sys
sys.path.insert(0, './app')
from dotenv import load_dotenv
load_dotenv('.env')

from db.database import db
from bson import ObjectId
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx

async def test():
    await db.connect()
    record = await db.db.resumes.find_one({})
    if not record:
        print("No records")
        return
        
    application_id = str(record["_id"])
    rd = record.get("resume_data", {})
    score = record.get("score", {})
    candidate_name = record.get("candidate_name", rd.get("name", "Candidate")) or "Candidate"
    candidate_email = record.get("candidate_email", rd.get("email", "")) or ""

    def safe(val):
        return str(val) if val is not None else ""

    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(safe(candidate_name))
    run.bold = True
    run.font.size = Pt(22)

    contact_parts = []
    if candidate_email:
        contact_parts.append(candidate_email)
    if rd.get("phone"):
        contact_parts.append(rd["phone"])
    if rd.get("location"):
        contact_parts.append(rd["location"])

    if contact_parts:
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s_run = subtitle.add_run(safe(" | ".join(contact_parts)))
        s_run.font.size = Pt(10)
        s_run.font.color.rgb = docx.shared.RGBColor(100, 116, 139)

    doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    total_score = score.get("total_score", 0) if isinstance(score, dict) else 0
    skill_score = score.get("skill_score", 0) if isinstance(score, dict) else 0
    exp_score   = score.get("experience_score", 0) if isinstance(score, dict) else 0

    score_p = doc.add_paragraph()
    score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = score_p.add_run(f"ATS Match Score: {total_score:.1f}%  |  ")
    r1.bold = True
    score_p.add_run(f"Skills: {skill_score:.1f}%  |  Experience: {exp_score:.1f}%")

    doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    def add_section(heading):
        h = doc.add_heading(heading, level=2)
        h.runs[0].font.color.rgb = docx.shared.RGBColor(30, 41, 59)

    summary = rd.get("summary", "") if isinstance(rd, dict) else ""
    if summary:
        add_section("Professional Summary")
        doc.add_paragraph(summary)

    skills = rd.get("skills", []) if isinstance(rd, dict) else []
    if skills:
        add_section("Skills")
        skills_text = " * ".join(skills) if isinstance(skills, list) else str(skills)
        doc.add_paragraph(skills_text)

    experience = rd.get("experience", []) if isinstance(rd, dict) else []
    if experience:
        add_section("Work Experience")
        if isinstance(experience, list):
            for exp in experience:
                if isinstance(exp, dict):
                    title_line = exp.get("title", "")
                    if exp.get("company"):
                        title_line += f"  --  {exp['company']}"
                    if exp.get("duration") or exp.get("years"):
                        title_line += f"  ({exp.get('duration') or exp.get('years', '')})"
                    p = doc.add_paragraph()
                    p.add_run(safe(title_line)).bold = True
                    if exp.get("description"):
                        doc.add_paragraph(str(exp["description"]), style="List Bullet")
                else:
                    doc.add_paragraph(str(exp), style="List Bullet")
        else:
            doc.add_paragraph(str(experience))

    print("Doc generation success for ID:", application_id)

try:
    asyncio.run(test())
except Exception as e:
    import traceback
    traceback.print_exc()
